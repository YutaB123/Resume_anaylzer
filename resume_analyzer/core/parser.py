"""Resume parsing module for PDF, DOCX, and TXT files."""

import io
from pathlib import Path
from typing import Union

from models.schemas import ResumeData, ResumeSection
from utils.helpers import clean_text


class ResumeParser:
    """Parser for extracting text from resume files."""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt'}
    
    def __init__(self):
        """Initialize the parser."""
        pass
    
    def parse(self, file_path: Union[str, Path] = None, file_bytes: bytes = None, 
              file_name: str = None) -> ResumeData:
        """Parse a resume file and extract text.
        
        Args:
            file_path: Path to the file (optional if file_bytes provided)
            file_bytes: Raw bytes of the file (for Gradio uploads)
            file_name: Name of the file (required if file_bytes provided)
            
        Returns:
            ResumeData object with extracted content
            
        Raises:
            ValueError: If file type is not supported
        """
        if file_path:
            file_path = Path(file_path)
            file_name = file_path.name
            extension = file_path.suffix.lower()
            
            with open(file_path, 'rb') as f:
                file_bytes = f.read()
        elif file_bytes and file_name:
            extension = Path(file_name).suffix.lower()
        else:
            raise ValueError("Either file_path or (file_bytes and file_name) must be provided")
        
        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {extension}. Supported: {self.SUPPORTED_EXTENSIONS}")
        
        # Extract text based on file type
        if extension == '.pdf':
            raw_text = self._parse_pdf(file_bytes)
        elif extension in {'.docx', '.doc'}:
            raw_text = self._parse_docx(file_bytes)
        elif extension == '.txt':
            raw_text = self._parse_txt(file_bytes)
        else:
            raw_text = ""
        
        # Clean the extracted text
        cleaned_text = clean_text(raw_text)
        
        return ResumeData(
            raw_text=cleaned_text,
            file_name=file_name,
            file_type=extension,
            sections=[]
        )
    
    def _parse_pdf(self, file_bytes: bytes) -> str:
        """Extract text from PDF file.
        
        Args:
            file_bytes: PDF file content as bytes
            
        Returns:
            Extracted text
        """
        try:
            import pdfplumber
            
            text_parts = []
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            return '\n\n'.join(text_parts)
        
        except ImportError:
            raise ImportError("pdfplumber is required for PDF parsing. Install with: pip install pdfplumber")
        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {str(e)}")
    
    def _parse_docx(self, file_bytes: bytes) -> str:
        """Extract text from DOCX file.
        
        Args:
            file_bytes: DOCX file content as bytes
            
        Returns:
            Extracted text
        """
        try:
            from docx import Document
            
            doc = Document(io.BytesIO(file_bytes))
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            return '\n'.join(text_parts)
        
        except ImportError:
            raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {str(e)}")
    
    def _parse_txt(self, file_bytes: bytes) -> str:
        """Extract text from TXT file.
        
        Args:
            file_bytes: TXT file content as bytes
            
        Returns:
            Extracted text
        """
        # Try common encodings
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                return file_bytes.decode(encoding)
            except UnicodeDecodeError:
                continue
        
        # Fallback with error handling
        return file_bytes.decode('utf-8', errors='replace')
    
    def get_file_info(self, file_bytes: bytes, file_name: str) -> dict:
        """Get basic file information without full parsing.
        
        Args:
            file_bytes: File content as bytes
            file_name: Name of the file
            
        Returns:
            Dictionary with file info
        """
        extension = Path(file_name).suffix.lower()
        
        return {
            "file_name": file_name,
            "file_type": extension,
            "file_size_kb": len(file_bytes) / 1024,
            "supported": extension in self.SUPPORTED_EXTENSIONS
        }
